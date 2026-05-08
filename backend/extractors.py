import requests
from bs4 import BeautifulSoup
import fitz  # PyMuPDF
import io
import docx

def extract_text_from_url(url: str) -> str:
    """Fetches a URL and extracts clean text from the HTML."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.extract()
            
        # Get text
        text = soup.get_text(separator=' ', strip=True)
        return text
    except Exception as e:
        raise Exception(f"Failed to extract text from URL: {str(e)}")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes."""
    try:
        # Open the PDF from bytes
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text() + "\n"
        pdf_document.close()
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes including paragraphs and tables."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
                    
        return '\n'.join(text)
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
